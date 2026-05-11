from langchain_core.documents import Document
import pdfplumber
import docx
from fastapi import UploadFile

def load_pdf(file: UploadFile):
    text = ""

    with pdfplumber.open(file.file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"

    return [
        Document(
            page_content=text,
            metadata={"source": "pdf"}
        )
    ]

def load_docx(file: UploadFile):
    doc = docx.Document(file.file)

    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return [
        Document(
            page_content=text,
            metadata={"source": "docx"}
        )
    ]